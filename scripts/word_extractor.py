#!/usr/bin/env python3
"""
Extract device compatibility from Word documents (.docx)
"""

from docx import Document
from docx.table import Table
import re
import os
from datetime import datetime
from typing import Dict, List, Any
from dataclasses import dataclass


@dataclass
class ExtractionResult:
    success: bool
    method: str
    devices: List[Dict[str, Any]]
    tables_found: int
    confidence: float
    errors: List[str]


class WordExtractor:
    """Extract devices from Word document tables."""
    
    def __init__(self, docx_path: str, region: str = "US"):
        self.docx_path = docx_path
        self.region = region
        
        self.KNOWN_MANUFACTURERS = [
            'Google', 'Samsung', 'OnePlus', 'Motorola', 'LG', 'HTC', 
            'Nokia', 'HMD Global', 'Sony', 'Xiaomi', 'Oppo', 'Vivo', 
            'Realme', 'Lively', 'TCL', 'Asus', 'Alcatel', 'Huawei', 'ZTE',
            'Apple'
        ]
    
    def extract(self) -> ExtractionResult:
        """Extract devices from Word document."""
        
        print(f"\n{'='*70}")
        print(f"🔍 EXTRACTING: {os.path.basename(self.docx_path)} ({self.region})")
        print(f"{'='*70}")
        
        devices = []
        errors = []
        tables_found = 0
        
        try:
            doc = Document(self.docx_path)
            
            print(f"\n📄 Document loaded: {len(doc.tables)} tables found")
            print(f"📝 Paragraphs: {len(doc.paragraphs)}\n")
            
            # Table mapping based on caption
            table_captions = {}
            
            # First pass: identify tables by their captions
            print("🔍 Scanning for table captions...")
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                
                # E3 iOS - Table 3
                if re.search(r'Table\s+3.*E3.*iOS', text, re.IGNORECASE):
                    table_captions[3] = ("E3", "ios")
                    print(f"   ✓ Table 3: E3 iOS")
                
                # E3 Android - Table 4
                elif re.search(r'Table\s+4.*E3.*Android', text, re.IGNORECASE):
                    table_captions[4] = ("E3", "android")
                    print(f"   ✓ Table 4: E3 Android")
                
                # 365 iOS - Table 5
                elif re.search(r'Table\s+5.*365.*iOS', text, re.IGNORECASE):
                    table_captions[5] = ("365", "ios")
                    print(f"   ✓ Table 5: 365 iOS")
                
                # 365 Android - Table 6
                elif re.search(r'Table\s+6.*365.*Android', text, re.IGNORECASE):
                    table_captions[6] = ("365", "android")
                    print(f"   ✓ Table 6: 365 Android")
                
                # NOW iOS - Table 7
                elif re.search(r'Table\s+7.*NOW.*iOS|Table\s+7.*iOS.*NOW', text, re.IGNORECASE):
                    table_captions[7] = ("NOW", "ios")
                    print(f"   ✓ Table 7: NOW iOS")
                
                # NOW Android - Table 8
                elif re.search(r'Table\s+8.*NOW.*Android|Table\s+8.*Android.*NOW', text, re.IGNORECASE):
                    table_captions[8] = ("NOW", "android")
                    print(f"   ✓ Table 8: NOW Android")
            
            print(f"\n✅ Found {len(table_captions)} table captions\n")
            
            # Second pass: extract from tables
            print("📋 Processing tables...")
            table_count = 0
            
            for table_idx, table in enumerate(doc.tables):
                table_count += 1
                
                # Determine table type
                table_info = None
                
                # Try to find which table this is by number
                if table_count in table_captions:
                    product, os_type = table_captions[table_count]
                    table_info = (product, os_type)
                    print(f"\n   Table {table_count}: {product} {os_type.upper()}")
                
                # If not found by number, try to infer from headers
                if not table_info and len(table.rows) > 0:
                    header_text = self._get_row_text(table.rows[0]).lower()
                    
                    if 'device manufacturer' in header_text:
                        print(f"\n   Table {table_count}: Device table (inferring type...)")
                        
                        # Try to find nearby caption
                        for table_num in sorted(table_captions.keys(), reverse=True):
                            if table_num <= table_idx:
                                table_info = table_captions[table_num]
                                break
                
                if not table_info:
                    print(f"   ⏭️  Table {table_count}: Skipped (not a device table)")
                    continue
                
                product, os_type = table_info
                
                # Show table info
                print(f"      📊 {len(table.rows)} rows, {len(table.rows[0].cells)} columns")
                
                if len(table.rows) > 0:
                    header = self._get_row_text(table.rows[0])
                    print(f"      Header: {header[:80]}")
                
                # Parse table
                parsed = self._parse_table(table, product, os_type)
                
                if parsed:
                    devices.extend(parsed)
                    tables_found += 1
                    print(f"      ✅ Extracted {len(parsed)} devices")
                else:
                    print(f"      ⚠️  No devices extracted")
            
            print(f"\n{'='*70}")
            print(f"✅ Total: {len(devices)} devices from {tables_found} tables")
            print(f"{'='*70}\n")
            
            return ExtractionResult(
                success=len(devices) > 0,
                method="word_docx",
                devices=devices,
                tables_found=tables_found,
                confidence=0.95 if len(devices) > 0 else 0.0,
                errors=errors
            )
            
        except Exception as e:
            errors.append(str(e))
            import traceback
            print(f"\n❌ Error: {e}")
            traceback.print_exc()
            return ExtractionResult(False, "failed", [], 0, 0.0, errors)
    
    
    def _get_row_text(self, row) -> str:
        """Extract text from a table row."""
        return ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
    
    
    def _parse_table(self, table: Table, product: str, os_type: str) -> List[Dict]:
        """Parse a device table."""
        
        devices = []
        
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:  # Skip header row
                continue
            
            cells = row.cells
            if len(cells) < 2:
                continue
            
            # Extract cell values and clean
            cell_values = []
            for cell in cells:
                text = cell.text.strip()
                if text:
                    cell_values.append(text)
            
            if len(cell_values) < 2:
                continue
            
            manufacturer = cell_values[0]
            model = cell_values[1]
            
            if not manufacturer or not model:
                continue
            
            # Skip if this is a header or metadata row
            if any(skip in manufacturer.lower() for skip in ['table', 'page', 'manufacturer', 'device']):
                continue
            
            # Parse based on OS
            if os_type == 'ios':
                # iOS devices
                if manufacturer.lower() != 'apple':
                    continue
                
                model_number = cell_values[2] if len(cell_values) > 2 else ''
                
                # Check for RQ (Rationally Qualified)
                rq = False
                if 'rationally' in model_number.lower() or 'qualified' in model_number.lower():
                    rq = True
                    model_number = ''
                
                # Handle concatenated models (e.g., "iPhone 11 iPhone 11 Pro")
                iphone_models = self._extract_iphone_models(model)
                
                for m in iphone_models:
                    if not m or len(m) < 3:
                        continue
                    
                    device = {
                        "name": f"Apple {m}",
                        "manufacturer": "Apple",
                        "model": m,
                        "model_number": model_number,
                        "os_version": self._get_ios_version(m),
                        "rationally_qualified": rq or product == "NOW",
                        "product": product,
                        "region": self.region
                    }
                    
                    devices.append(device)
            
            else:
                # Android devices
                if manufacturer not in self.KNOWN_MANUFACTURERS:
                    continue
                
                model_number = ''
                rq = False
                
                # Extract RQ status from remaining cells
                if len(cell_values) > 2:
                    third_col = cell_values[2]
                    
                    # Check if it's RQ status
                    if third_col.lower() == 'yes':
                        rq = True
                    elif 'no (' in third_col.lower():
                        # Extract model number from "No (GD1YQ)"
                        match = re.search(r'No\s*\(([^)]+)\)', third_col, re.IGNORECASE)
                        if match:
                            model_number = match.group(1)
                    else:
                        model_number = third_col
                
                # Check fourth column for RQ if we haven't found it
                if not rq and len(cell_values) > 3:
                    fourth_col = cell_values[3]
                    if fourth_col.lower() == 'yes':
                        rq = True
                
                # Skip if model is too short
                if len(model) < 2:
                    continue
                
                device = {
                    "name": f"{manufacturer} {model}",
                    "manufacturer": manufacturer,
                    "model": model,
                    "model_number": model_number,
                    "os_version": self._get_android_version(model),
                    "rationally_qualified": rq,
                    "product": product,
                    "region": self.region
                }
                
                devices.append(device)
        
        return devices
    
    
    def _extract_iphone_models(self, text: str) -> List[str]:
        """Extract iPhone model names from text."""
        
        # Pattern matches: iPhone 14, iPhone 14 Pro, iPhone 14 Pro Max, iPhone SE, iPhone X, etc.
        pattern = r'iPhone\s+(?:\d+\s*(?:Pro\s*(?:Max)?|Plus|Mini)?|SE(?:\s+\(\d+(?:st|nd|rd|th)\s+generation\))?|X[RS]?(?:\s+Max)?)'
        
        matches = re.finditer(pattern, text, re.IGNORECASE)
        models = []
        seen = set()
        
        for match in matches:
            model = match.group(0).strip()
            # Normalize spaces
            model = ' '.join(model.split())
            
            # Skip duplicates and noise
            if model in seen or model.lower() in ['confidential', 'rationally', 'qualified']:
                continue
            
            seen.add(model)
            models.append(model)
        
        # Fallback: if no matches found, return the text as-is
        if not models and text and len(text.strip()) > 2:
            models = [text.strip()]
        
        return models
    
    
    def _get_ios_version(self, model: str) -> str:
        """Get iOS version for model."""
        
        m = model.lower()
        
        # Map model to iOS version
        if 'iphone 16' in m:
            return '18.0'
        elif 'iphone 15' in m:
            return '17.0'
        elif 'iphone 14' in m:
            return '16.0'
        elif 'iphone 13' in m:
            return '15.0'
        elif 'iphone 12' in m:
            return '14.0'
        elif 'iphone 11' in m:
            return '13.0'
        elif 'iphone x' in m or 'iphone xs' in m or 'iphone xr' in m:
            return '12.0'
        elif 'iphone 8' in m:
            return '11.0'
        elif 'watch' in m:
            return '9.0'
        
        # Default
        return '12.0'
    
    
    def _get_android_version(self, model: str) -> str:
        """Get Android version for model."""
        
        m = model.lower()
        
        # Map model to Android version
        if 'pixel 9' in m or 'pixel 8' in m:
            return '14.0'
        elif 'pixel 7' in m or 'pixel 6' in m or 'pixel 5' in m:
            return '13.0'
        elif 'pixel 4' in m:
            return '12.0'
        elif 's25' in m or 's24' in m:
            return '15.0'
        elif 's23' in m or 's22' in m:
            return '14.0'
        
        # Default
        return '10.0'


def parse_compatibility_docx(docx_path: str, region: str = "US") -> Dict[str, Any]:
    """
    Parse Word document and return compatibility data.
    
    Args:
        docx_path: Path to .docx file
        region: Region code (US or OUS)
    
    Returns:
        Dictionary with products and device data
    """
    
    extractor = WordExtractor(docx_path, region)
    result = extractor.extract()
    
    compatibility_data = {
        "region": region,
        "last_updated": datetime.now().isoformat(),
        "document_info": {
            "source_file": os.path.basename(docx_path),
            "extraction_method": result.method,
            "confidence": result.confidence
        },
        "products": {
            "E3": {"android": [], "ios": []},
            "365": {"android": [], "ios": []},
            "NOW": {"android": [], "ios": []}
        }
    }
    
    # Organize devices by product and OS
    for device in result.devices:
        product = device.get('product', 'E3')
        os_type = device.get('os', 'android')
        
        # Validate product and OS
        if product not in compatibility_data["products"]:
            continue
        
        if os_type not in compatibility_data["products"][product]:
            continue
        
        # Check for duplicates
        existing_names = [
            d['name'] 
            for d in compatibility_data["products"][product][os_type]
        ]
        
        if device['name'] not in existing_names:
            compatibility_data["products"][product][os_type].append(device)
    
    return compatibility_data


# For direct testing
if __name__ == '__main__':
    import json
    
    # Test with US document
    if os.path.exists('docx/compatibility_us.docx'):
        print("\n🧪 Testing US document extraction...")
        us_data = parse_compatibility_docx('docx/compatibility_us.docx', region="US")
        
        # Print summary
        total = 0
        for product in ['E3', '365', 'NOW']:
            ios = len(us_data['products'][product]['ios'])
            android = len(us_data['products'][product]['android'])
            product_total = ios + android
            total += product_total
            print(f"{product}: {ios} iOS + {android} Android = {product_total}")
        
        print(f"\nTotal: {total} devices")
        
        # Save sample
        with open('sample_us.json', 'w') as f:
            json.dump(us_data, f, indent=2)
        print("\nSample saved to: sample_us.json")
    else:
        print("❌ docx/compatibility_us.docx not found")
